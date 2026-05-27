import os
import sys
import time
import argparse
import json
import tempfile
import hashlib
from pathlib import Path
import google.generativeai as genai
from rich.console import Console
from pygments import highlight
from pygments.lexers import get_lexer_by_name
from pygments.formatters import ImageFormatter

console = Console()

MODELS = {
    "flash": "gemini-1.5-flash",
    "pro": "gemini-1.5-pro"
}

def renderizar_codigo_a_imagen(codigo, filename, output_path, theme="monokai"):
    try:
        lexer = get_lexer_by_name("php", stripall=True)
        formatter = ImageFormatter(font_name="DejaVu Sans Mono", font_size=20, line_numbers=True, style=theme)
        with open(output_path, "wb") as f:
            f.write(highlight(codigo, lexer, formatter))
        return True
    except Exception as e:
        console.print(f"[red]Error al renderizar imagen: {e}[/red]")
        return False

def append_to_docx(doc_obj, text):
    text = "".join(c for c in text if c.isprintable() or c in "\n\r\t")
    lineas = text.split('\n')
    for linea in lineas:
        if linea.startswith('# '):
            doc_obj.add_heading(linea[2:], level=1)
        elif linea.startswith('## '):
            doc_obj.add_heading(linea[3:], level=2)
        elif linea.startswith('### '):
            doc_obj.add_heading(linea[4:], level=3)
        elif linea.startswith('- ') or linea.startswith('* '):
            doc_obj.add_paragraph(linea[2:], style='List Bullet')
        else:
            doc_obj.add_paragraph(linea)

def generar_doc_yield(path, modelo_alias, formato='md', provider='gemini', api_key=None, template=None, custom_prompt=None, resume=True, font_name="Calibri", font_size=11, block_size=100, diff_mode=False, structure_only=False, cancel_token=None, code_theme="monokai", exclude_patterns=""):
    if not os.path.exists(path) or not os.path.isdir(path):
        yield {"error": f"La ruta '{path}' no existe o no es un directorio", "fatal": True}
        return
    start_time = time.time()
    nombre_proyecto = Path(path).resolve().name
    
    # 1. Escaneo Inicial
    archivos_totales = []
    ignorar_dirs = ['vendor', 'node_modules', 'storage', 'public', 'tests', 'database', 'dist', 'build', '.git']
    ignorar_files = ['webpack.mix.js', 'tailwind.config.js', 'package-lock.json']
    excluir_patrones = [p.strip() for p in exclude_patterns.split(",") if p.strip()] if exclude_patterns else []
    
    for root, dirs, files in os.walk(path):
        dirs[:] = [d for d in dirs if d not in ignorar_dirs and not d.startswith('.')]
        for file in files:
            # Filtro por extensión
            if file.endswith(('.php', '.js', '.vue', '.ts', '.tsx', '.blade.php', '.css', '.scss', '.py', '.go', '.rb', '.java', '.rs', '.kt', '.swift')):
                if file.endswith('.min.js') or file in ignorar_files:
                    continue
                
                full_p = Path(root) / file
                rel_path = os.path.relpath(full_p, path)
                
                # Filtro de exclusión personalizado
                if excluir_patrones:
                    ruta_lower = rel_path.replace("\\", "/").lower()
                    if any(
                        ruta_lower == p.lower() or ruta_lower.startswith(p.lower() + "/")
                        for p in excluir_patrones
                    ):
                        continue
                
                try:
                    with open(full_p, "r", encoding="utf-8", errors="ignore") as f:
                        archivos_totales.append((rel_path, full_p))
                except Exception: continue

    total_bloques = 0
    total_lineas = 0
    for _, full_p in archivos_totales:
        try:
            with open(full_p, "r", encoding="utf-8", errors="ignore") as f:
                nlines = len(f.readlines())
                total_bloques += (nlines // block_size) + 1
                total_lineas += nlines
        except Exception: continue

    yield {"status": "Iniciando...", "progress": 5, "total_bloques": total_bloques, "total_lineas": total_lineas}

    # Modo solo estructura: árbol de proyecto sin AI
    if structure_only:
        Path("docs_laravel").mkdir(parents=True, exist_ok=True)
        out_file = f"docs_laravel/estructura_{nombre_proyecto.lower()}.md"
        with open(out_file, "w", encoding="utf-8") as f:
            f.write(f"# Estructura del Proyecto: {nombre_proyecto}\n\n")
            f.write(f"Generado: {time.strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            f.write("## Árbol de Archivos\n\n")
            tree = {}
            for rel_path, full_p in archivos_totales:
                parts = rel_path.split(os.sep)
                node = tree
                for p in parts[:-1]:
                    node = node.setdefault(p, {})
                node[parts[-1]] = None
            def escribir_arbol(nodo, prefijo="", es_ultimo=True):
                items = list(nodo.items())
                for i, (nombre, sub) in enumerate(items):
                    es_ult = i == len(items) - 1
                    conector = "└── " if es_ult else "├── "
                    f.write(f"{prefijo}{conector}{nombre}\n")
                    if sub is not None:
                        extension = "    " if es_ult else "│   "
                        escribir_arbol(sub, prefijo + extension, es_ult)
            escribir_arbol(tree)
            f.write("\n## Resumen\n\n")
            exts = {}
            for rel_path, _ in archivos_totales:
                ext = rel_path.split(".")[-1] if "." in rel_path else "sin_ext"
                exts[ext] = exts.get(ext, 0) + 1
            f.write(f"- **Total archivos**: {len(archivos_totales)}\n")
            f.write(f"- **Extensiones**:\n")
            for ext, count in sorted(exts.items(), key=lambda x: -x[1]):
                f.write(f"  - `.{ext}`: {count} archivos\n")
            if total_lineas:
                f.write(f"- **Líneas totales**: ~{total_lineas}\n")
        yield {"status": "Completado (Estructura)", "progress": 100, "file": out_file, "summary": {"archivos": len(archivos_totales), "bloques": 0, "segundos": int(time.time() - start_time)}}
        return

    # Modo diff: cargar manifest y filtrar archivos sin cambios
    manifest_file = Path(path) / ".auditor_manifest.json"
    manifest = {}
    archivos_a_procesar = archivos_totales.copy()
    if diff_mode and manifest_file.exists():
        try:
            with open(manifest_file, "r") as f:
                manifest = json.load(f)
            filtrados = []
            for rel_path, full_p in archivos_totales:
                mtime_actual = os.path.getmtime(full_p)
                if rel_path in manifest.get("files", {}) and manifest["files"][rel_path] == mtime_actual:
                    continue
                filtrados.append((rel_path, full_p))
            saltados = len(archivos_totales) - len(filtrados)
            archivos_a_procesar = filtrados
            if saltados:
                yield {"log": f"📋 Modo diff: {saltados} archivos sin cambios, {len(filtrados)} por procesar"}
        except Exception:
            yield {"log": "⚠️ Manifest corrupto, procesando todos los archivos"}

    # Cache de respuestas AI
    cache_file = Path(path) / ".auditor_cache.json"
    cache = {}
    if cache_file.exists():
        try:
            with open(cache_file, "r") as f:
                cache = json.load(f)
            yield {"log": f"💾 Cache cargado ({len(cache)} entradas)"}
        except Exception:
            yield {"log": "⚠️ Cache corrupto, iniciando fresco"}
    
    # Checkpoint
    checkpoint_file = Path(path) / ".auditor_state.json"
    last_state = None
    if resume and checkpoint_file.exists():
        try:
            with open(checkpoint_file, "r") as f:
                last_state = json.load(f)
            yield {"log": f"🔄 Reanudando desde {last_state.get('idx_arch', 0)}"}
        except Exception:
            yield {"log": "⚠️ Checkpoint corrupto, iniciando desde cero"}
            last_state = None

    # Si no hay checkpoint pero existe el archivo de salida, escanearlo para reanudar
    if not last_state and resume and os.path.exists(out_file):
        try:
            if formato == 'docx':
                from docx import Document
                doc_existente = Document(out_file)
                bloques_encontrados = []
                for parrafo in doc_existente.paragraphs:
                    if parrafo.style.name.startswith('Heading') and 'Archivo: ' in parrafo.text:
                        # "Archivo: app/Models/User.php (Bloque 2)"
                        resto = parrafo.text.replace('Archivo: ', '')
                        if ' (Bloque ' in resto:
                            parts = resto.split(' (Bloque ')
                            archivo = parts[0]
                            bloque = parts[1].rstrip(')')
                            bloques_encontrados.append((archivo, int(bloque)))
                if bloques_encontrados:
                    ultimo_archivo = bloques_encontrados[-1][0]
                    ultimo_bloque = bloques_encontrados[-1][1]
                    for idx, (rel, _) in enumerate(archivos_a_procesar):
                        if rel == ultimo_archivo:
                            last_state = {"idx_arch": idx, "block_idx": (ultimo_bloque - 1) * block_size}
                            yield {"log": f"🔄 Reanudando DOCX existente: {ultimo_archivo} B{ultimo_bloque} ({len(bloques_encontrados)} bloques)"}
                            break
            else:
                import re
                with open(out_file, "r", encoding="utf-8") as f:
                    contenido_existente = f.read()
                bloques_encontrados = re.findall(r'^## (.+?) - Bloque (\d+)', contenido_existente, re.MULTILINE)
                if bloques_encontrados:
                    ultimo_archivo = bloques_encontrados[-1][0]
                    ultimo_bloque = int(bloques_encontrados[-1][1])
                    for idx, (rel, _) in enumerate(archivos_a_procesar):
                        if rel == ultimo_archivo:
                            last_state = {"idx_arch": idx, "block_idx": (ultimo_bloque - 1) * block_size}
                            yield {"log": f"🔄 Reanudando desde archivo existente: {ultimo_archivo} B{ultimo_bloque} ({len(bloques_encontrados)} bloques)"}
                            break
        except Exception as e:
            yield {"log": f"⚠️ No se pudo escanear archivo existente para reanudar: {e}"}

    # Configurar Clientes
    if provider == 'openrouter':
        from openai import OpenAI
        client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=api_key)
    else:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(MODELS.get(modelo_alias, modelo_alias))

    Path("docs_laravel").mkdir(parents=True, exist_ok=True)
    out_file = f"docs_laravel/auditoria_{nombre_proyecto.lower()}.{formato}"
    doc_obj = None
    if formato == 'docx':
        from docx import Document
        if last_state and os.path.exists(out_file):
            doc_obj = Document(out_file)
            yield {"log": f"📂 Cargando DOCX existente para continuar"}
        else:
            doc_obj = Document(Path("templates") / template) if template and (Path("templates") / template).exists() else Document()
        from docx.shared import Pt
        style = doc_obj.styles['Normal']
        style.font.name = font_name
        style.font.size = Pt(font_size)
        for level in range(1, 4):
            style = doc_obj.styles[f'Heading {level}']
            style.font.name = font_name
    
    bloques_procesados = 0
    archivos_procesados = 0

    # Limpiar output si es una corrida fresh (no resume)
    if not last_state and formato != 'docx':
        with open(out_file, "w", encoding="utf-8") as f:
            pass

    # 2. Bucle Principal
    for idx_arch, (rel_path, full_path) in enumerate(archivos_a_procesar):
        if cancel_token and cancel_token.is_set():
            yield {"log": "🛑 Auditoría cancelada por el usuario", "fatal": True}
            return
        if last_state and idx_arch < last_state.get('idx_arch', 0):
            try:
                with open(full_path, "r", encoding="utf-8", errors="ignore") as f:
                    bloques_procesados += (len(f.readlines()) // block_size) + 1
            except Exception: pass
            continue

        # Calcular bloques totales de este archivo
        try:
            with open(full_path, "r", encoding="utf-8", errors="ignore") as f:
                lineas_file = f.readlines()
            total_file_blocks = (len(lineas_file) // block_size) + 1
        except Exception:
            continue

        yield {
            "log": f"📖 Analizando {rel_path}...",
            "file": {"name": rel_path, "blocks_total": total_file_blocks, "blocks_done": 0, "total_files": len(archivos_a_procesar), "file_index": idx_arch}
        }

        try:
            with open(full_path, "r", encoding="utf-8", errors="ignore") as f:
                lineas = f.readlines()
        except Exception: continue

        for i in range(0, len(lineas), block_size):
            if last_state and idx_arch == last_state.get('idx_arch', 0) and i <= last_state.get('block_idx', -1):
                bloques_procesados += 1
                continue

            bloque_codigo = "".join(lineas[i:i+block_size])
            num_bloque = (i // block_size) + 1
            bloques_procesados += 1
            
            # 📸 CAPTURA DE CÓDIGO (Solo para DOCX)
            img_path = None
            if formato == 'docx':
                img_path = Path(tempfile.gettempdir()) / f"audit_{idx_arch}_{num_bloque}.png"
                renderizar_codigo_a_imagen(bloque_codigo, rel_path, img_path, code_theme)

            try:
                if cancel_token and cancel_token.is_set():
                    yield {"error": "🛑 Auditoría cancelada por el usuario", "fatal": True}
                    return
                cache_key = hashlib.md5(f"{rel_path}:{i}:{bloque_codigo}".encode()).hexdigest()
                if cache_key in cache:
                    texto_explicacion = cache[cache_key]["respuesta"]
                    tokens_usados = cache[cache_key].get("tokens", 0)
                    exito = True
                    yield {"log": f"💾 Cache hit para {rel_path} B{num_bloque} ({len(bloque_codigo)} chars)", "tokens": tokens_usados}
                else:
                    intentos = 0
                    exito = False
                    texto_explicacion = ""
                    tokens_usados = None
                while intentos < 3 and not exito:
                    try:
                        p = custom_prompt.replace("{{CODE}}", bloque_codigo).replace("{{PROJECT}}", nombre_proyecto).replace("{{FILE}}", rel_path).replace("{{RANGE}}", f"{i+1}-{i+len(lineas[i:i+block_size])}") if custom_prompt else f"Analiza: {bloque_codigo}"
                        
                        if provider == 'openrouter':
                            res = client.chat.completions.create(
                                model=modelo_alias, 
                                messages=[{"role": "user", "content": p}],
                                timeout=60
                            )
                            if res and res.choices and len(res.choices) > 0:
                                texto_explicacion = res.choices[0].message.content or "Sin respuesta del modelo."
                                exito = True
                                try:
                                    if res.usage:
                                        tokens_usados = res.usage.total_tokens
                                except: pass
                            else:
                                raise Exception("OpenRouter devolvió una respuesta vacía.")
                        else:
                            res = model.generate_content(p)
                            if res and res.text:
                                texto_explicacion = res.text
                                exito = True
                                try:
                                    if hasattr(res, 'usage_metadata') and res.usage_metadata:
                                        tokens_usados = res.usage_metadata.total_token_count or 0
                                except: pass
                            else:
                                raise Exception("Gemini devolvió una respuesta vacía.")
                                
                    except Exception as e:
                        intentos += 1
                        err_msg = str(e)
                        if hasattr(e, 'body') and isinstance(e.body, dict):
                            err_body = e.body.get('error', {})
                            if isinstance(err_body, dict):
                                raw_msg = err_body.get('message', '')
                                status = getattr(e, 'status_code', None) or err_body.get('code', '')
                                if status:
                                    err_msg = f"[{status}] {raw_msg}" if raw_msg else str(e)
                                elif raw_msg:
                                    err_msg = raw_msg
                        yield {"log": f"⚠️ Intento {intentos}/3: {err_msg}", "tokens": tokens_usados}
                        time.sleep(2 ** intentos)

                if not exito:
                    yield {"error": "Error fatal después de 3 intentos. Revisa tu conexión y límites de API.", "fatal": True}
                    return

                # Guardar
                if formato == 'docx':
                    from docx.shared import Inches
                    doc_obj.add_heading(f"Archivo: {rel_path} (Bloque {num_bloque})", level=2)
                    if img_path and img_path.exists():
                        doc_obj.add_picture(str(img_path), width=Inches(6))
                    append_to_docx(doc_obj, texto_explicacion)
                    doc_obj.save(out_file)
                else:
                    bloque_output = f"## {rel_path} - Bloque {num_bloque}\n\n{texto_explicacion}\n\n"
                    with open(out_file, "a", encoding="utf-8") as f:
                        f.write(bloque_output)

                # Guardar en cache (solo si no vino de cache)
                if cache_key not in cache:
                    cache[cache_key] = {"respuesta": texto_explicacion, "tokens": tokens_usados}
                    with open(cache_file, "w") as f:
                        json.dump(cache, f)

                # Checkpoint
                with open(checkpoint_file, "w") as f:
                    json.dump({"idx_arch": idx_arch, "block_idx": i}, f)
                
                yield {
                    "progress": int(10 + (bloques_procesados/total_bloques)*88),
                    "log": f"✅ {rel_path} B{num_bloque} OK",
                    "file": {"name": rel_path, "blocks_done": num_bloque},
                    "tokens": tokens_usados,
                    "live_content": bloque_output if formato != 'docx' else None
                }
            finally:
                if img_path and img_path.exists():
                    img_path.unlink()
        
        archivos_procesados += 1

    # Limpiar checkpoint al finalizar
    if checkpoint_file.exists(): checkpoint_file.unlink()

    # Guardar manifest para modo diff
    if diff_mode:
        nuevo_manifest = {"files": {}}
        for rel_path, full_p in archivos_totales:
            try:
                nuevo_manifest["files"][rel_path] = os.path.getmtime(full_p)
            except: pass
        with open(manifest_file, "w") as f:
            json.dump(nuevo_manifest, f)

    archivos_saltados = len(archivos_totales) - len(archivos_a_procesar)
    tiempo_total = int(time.time() - start_time)

    # Conversión a PDF si se solicitó
    if formato == 'pdf' and os.path.exists(out_file):
        try:
            from weasyprint import HTML
            pdf_file = out_file.replace('.md', '.pdf')
            html_content = "<html><head><meta charset='utf-8'><style>body{font-family:DejaVu Sans, sans-serif;font-size:11pt;line-height:1.6;max-width:800px;margin:40px auto;color:#333;}h1{color:#1a1a2e;border-bottom:2px solid #1a1a2e;padding-bottom:8px;}h2{color:#16213e;margin-top:32px;}h3{color:#0f3460;}code{background:#f4f4f4;padding:2px 6px;border-radius:4px;font-size:10pt;}pre{background:#f8f8f8;padding:16px;border-radius:8px;overflow-x:auto;}img{max-width:100%;}</style></head><body>"
            with open(out_file, "r", encoding="utf-8") as f:
                md_text = f.read()
            import re
            html_body = re.sub(r'^### (.+)$', r'<h3>\1</h3>', md_text, flags=re.MULTILINE)
            html_body = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html_body, flags=re.MULTILINE)
            html_body = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html_body, flags=re.MULTILINE)
            html_body = re.sub(r'```(\w*)\n(.*?)```', r'<pre><code>\2</code></pre>', html_body, flags=re.DOTALL)
            html_body = re.sub(r'`([^`]+)`', r'<code>\1</code>', html_body)
            html_body = html_body.replace('\n', '<br>')
            html_content += html_body + "</body></html>"
            HTML(string=html_content).write_pdf(pdf_file)
            out_file = pdf_file
            yield {"log": f"📄 PDF generado: {pdf_file}"}
        except Exception as e:
            yield {"log": f"⚠️ Error generando PDF: {e}"}

    yield {
        "status": "Completado",
        "progress": 100,
        "file": out_file,
        "summary": {
            "archivos": archivos_procesados,
            "archivos_totales": len(archivos_totales),
            "archivos_saltados": archivos_saltados,
            "bloques": bloques_procesados,
            "segundos": tiempo_total
        }
    }

    # Notificación desktop
    try:
        import subprocess
        subprocess.run(["notify-send", "Laravel Auditor PRO",
            f"Auditoría de {nombre_proyecto} completada en {tiempo_total}s\n{archivos_procesados} archivos, {bloques_procesados} bloques",
            "-i", "dialog-information", "-t", "5000"], timeout=5)
    except Exception:
        pass

def generar_doc_proyecto(path, modelo_alias, formato='md', provider='gemini', api_key=None):
    for update in generar_doc_yield(path, modelo_alias, formato, provider=provider, api_key=api_key):
        if "log" in update: print(update['log'])

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("path")
    args = parser.parse_args()
    generar_doc_proyecto(args.path, "pro")